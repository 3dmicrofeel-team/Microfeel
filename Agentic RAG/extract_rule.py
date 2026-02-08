#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""提取Rule.docx内容"""

try:
    from docx import Document
    import json
    import sys
    
    doc = Document('Rule.docx')
    
    # 提取所有段落
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 提取表格
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # 只保存非空行
                table_data.append(row_data)
        if table_data:
            tables_data.append(table_data)
    
    # 保存为JSON
    output = {
        'paragraphs': paragraphs,
        'tables': tables_data
    }
    
    # 输出到文件
    with open('rule_extracted.json', 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    
    # 同时输出文本版本
    with open('rule_extracted.txt', 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(paragraphs))
        if tables_data:
            f.write('\n\n=== 表格内容 ===\n\n')
            for i, table in enumerate(tables_data, 1):
                f.write(f'\n表格 {i}:\n')
                for row in table:
                    f.write(' | '.join(row) + '\n')
    
    print(f"成功提取 {len(paragraphs)} 个段落和 {len(tables_data)} 个表格")
    print("内容已保存到 rule_extracted.json 和 rule_extracted.txt")
    
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)
except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
